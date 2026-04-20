import os
import io
import re
import asyncio
import tempfile
import logging
from fastapi import UploadFile, HTTPException
from pdf2docx import Converter
from docx import Document
from deep_translator import GoogleTranslator

# Re-use the LibreOffice converter logic from word_to_pdf_service.py to turn the new DOCX back into a PDF
from services.word_to_pdf_service import _find_libreoffice, _run_libreoffice

logger = logging.getLogger(__name__)

# Characters that are typically not translatable (numbers, symbols, punctuation only)
_NON_TRANSLATABLE_RE = re.compile(r'^[\d\s\W_]+$')

def _is_skippable(text: str) -> bool:
    """Returns True if text has no real words worth translating."""
    stripped = text.strip()
    if not stripped or len(stripped) <= 1:
        return True
    if _NON_TRANSLATABLE_RE.match(stripped):
        return True
    return False

def _safe_translate(translator: GoogleTranslator, text: str, timeout: int = 15) -> str | None:
    """
    Translate a single string with a hard timeout to prevent hangs.
    Returns the translated string or None if it fails/times out.
    """
    import concurrent.futures
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(translator.translate, text)
        try:
            return future.result(timeout=timeout)
        except concurrent.futures.TimeoutError:
            logger.warning(f"Translation timed out for text: {text[:50]}...")
            return None
        except Exception as e:
            logger.warning(f"Translation error: {e}")
            return None

async def translate_pdf(file: UploadFile, target_lang: str):
    """
    Translates a PDF document by:
    1. Converting it to a Word (.docx) document using pdf2docx.
    2. Translating the text inside the Word document using python-docx and deep_translator.
    3. Converting the translated Word document back to a PDF using LibreOffice.
    """
    if not file.filename or not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are accepted.")

    logger.info(f"Translating PDF to {target_lang}: {file.filename}")

    # Read uploaded PDF bytes
    pdf_bytes = await file.read()
    if len(pdf_bytes) == 0:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    # Locate LibreOffice early to fail fast if it's missing
    try:
        soffice = _find_libreoffice()
    except RuntimeError as e:
        logger.error(str(e))
        raise HTTPException(status_code=503, detail=str(e))

    with tempfile.TemporaryDirectory() as temp_dir:
        input_pdf_path = os.path.join(temp_dir, "input.pdf")
        temp_docx_path = os.path.join(temp_dir, "temp_layout.docx")
        translated_docx_path = os.path.join(temp_dir, "translated.docx")
        final_pdf_path = os.path.join(temp_dir, "translated.pdf")

        # 1. Write PDF to temp file
        with open(input_pdf_path, "wb") as f:
            f.write(pdf_bytes)

        # 2. Convert PDF to DOCX
        try:
            logger.info("Converting PDF to DOCX (Step 1/3)...")
            cv = Converter(input_pdf_path)
            # Layout settings that improve fidelity:
            # - connected_border_tolerance: merges nearby borders so table cells stay together
            # - line_overlap_threshold: reduces false multi-line splits
            # - line_break_width_ratio: helps with line spacing accuracy
            # - lines_left_aligned_threshold: avoids falsely splitting columns
            cv.convert(
                temp_docx_path,
                start=0,
                end=None,
                connected_border_tolerance=0.5,
                line_overlap_threshold=0.9,
                line_break_width_ratio=0.5,
                lines_left_aligned_threshold=0.1,
            )
            cv.close()
        except Exception as e:
            logger.error(f"PDF to DOCX conversion failed: {e}")
            raise HTTPException(
                status_code=500, 
                detail=f"Layout Extraction Failed: We couldn't parse the PDF structure. " 
                       f"This usually happens with very complex or scanned PDFs. Error: {str(e)}"
            )

        # 3. Translate the DOCX
        def do_translate():
            logger.info("Translating content (Step 2/3)...")
            # Initialize translator
            try:
                translator = GoogleTranslator(source='auto', target=target_lang)
            except Exception as e:
                logger.error(f"Failed to initialize translator: {e}")
                raise HTTPException(status_code=400, detail=f"Unsupported language code: {target_lang}")

            if not os.path.exists(temp_docx_path):
                 raise RuntimeError("Intermediate DOCX file was not found before translation.")

            doc = Document(temp_docx_path)
            
            # Gather all runs that have text
            valid_runs = []
            
            for p in doc.paragraphs:
                for run in p.runs:
                    if run.text and not _is_skippable(run.text):
                        valid_runs.append(run)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if run.text and not _is_skippable(run.text):
                                    valid_runs.append(run)

            if not valid_runs:
                doc.save(translated_docx_path)
                return

            delimiter = "\n\n"
            max_chars = 4500
            current_chunk = []
            current_length = 0
            chunks = []

            for run in valid_runs:
                text = run.text.strip()
                if not text:
                    continue
                # If adding this text exceeds max_chars, finalize the chunk
                if current_length + len(text) + len(delimiter) > max_chars and current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = []
                    current_length = 0
                
                current_chunk.append((run, run.text, text))
                current_length += len(text) + len(delimiter)

            if current_chunk:
                chunks.append(current_chunk)

            logger.info(f"Translating {len(valid_runs)} runs in {len(chunks)} chunks...")

            for item_chunk in chunks:
                combined_text = delimiter.join([item[2] for item in item_chunk])
                try:
                    translated_combined = translator.translate(combined_text)
                    if not translated_combined:
                        continue
                    
                    # Normalize translated newlines occasionally warped by the translation service
                    normalized_translation = translated_combined.replace('\r\n', '\n').replace('\n \n', '\n\n')
                    split_res = [s.strip() for s in normalized_translation.split(delimiter)]

                    # Fallback if split mismatch is drastic
                    if len(split_res) != len(item_chunk):
                        logger.warning(f"Chunk split mismatch! Expected {len(item_chunk)}, got {len(split_res)}. Fallback to sequential for this chunk.")
                        # Attempt sequential fallback with individual timeouts
                        for run, original_text, stripped in item_chunk:
                            t_val = _safe_translate(translator, stripped)
                            if t_val:
                                out = t_val.strip()
                                if original_text.startswith(' '): out = ' ' + out
                                if original_text.startswith('  '): out = ' ' + out
                                if original_text.endswith(' '): out = out + ' '
                                run.text = out
                    else:
                        # Success, assign mapped values
                        for i, (run, original_text, stripped) in enumerate(item_chunk):
                            t_val = split_res[i]
                            if t_val:
                                out = t_val.strip()
                                if original_text.startswith(' '): out = ' ' + out
                                if original_text.startswith('  '): out = ' ' + out
                                if original_text.endswith(' '): out = out + ' '
                                run.text = out

                except Exception as e:
                    logger.warning(f"Translation failed for a chunk: {e}. Fallback to sequential.")
                    # fallback to sequential translation if batch fails totally, each with a timeout
                    for run, original_text, stripped in item_chunk:
                        t_val = _safe_translate(translator, stripped)
                        if t_val:
                            out = t_val.strip()
                            if original_text.startswith(' '): out = ' ' + out
                            if original_text.startswith('  '): out = ' ' + out
                            if original_text.endswith(' '): out = out + ' '
                            run.text = out

            # Let's save the file
            doc.save(translated_docx_path)

        # Run translation in thread pool
        try:
            await asyncio.to_thread(do_translate)
        except HTTPException as he:
            raise he
        except Exception as e:
            logger.error(f"Text translation failed (Step 2/3): {e}")
            raise HTTPException(
                status_code=500, 
                detail=f"Cloud Translation Failed: We couldn't translate the text content. "
                       f"Please try again in a few moments. Error: {str(e)}"
            )

        # 4. Convert translated DOCX back to PDF
        try:
            logger.info("Generating Final PDF (Step 3/3)...")
            if not os.path.exists(translated_docx_path):
                 raise RuntimeError("Translated DOCX file was not found.")
                 
            await asyncio.to_thread(_run_libreoffice, soffice, translated_docx_path, temp_dir)
        except Exception as e:
            logger.error(f"Final PDF generation failed (Step 3/3): {e}")
            raise HTTPException(
                status_code=500, 
                detail=f"PDF Generation Failed: We translated the text but couldn't rebuild the PDF file. "
                       f"Error: {str(e)}"
            )

        if not os.path.isfile(final_pdf_path):
            logger.error(f"Expected PDF at {final_pdf_path} not found.")
            raise HTTPException(
                status_code=500, 
                detail="System Error: The translated PDF was not generated correctly by the backend."
            )

        # Read the final PDF into memory
        with open(final_pdf_path, "rb") as f:
            pdf_bytes = f.read()

        # Build the output filename
        base_name = os.path.splitext(file.filename)[0]
        output_filename = f"{base_name}_{target_lang}.pdf"

    return pdf_bytes, output_filename
